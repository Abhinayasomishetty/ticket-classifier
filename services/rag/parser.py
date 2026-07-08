# Document parsing utilities.
# Supports PDF, DOCX, and TXT files.
import os
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path
from llama_index.core import Document as LlamaDocument
import fitz  # PyMuPDF
import docx  # python-docx


def parse_pdf(file_path: str) -> Tuple[List[str], Dict[str, Any]]:
    """
    Parse a PDF file into text chunks (one per page).

    Returns:
        Tuple of (list of text chunks, metadata dict)
    """
    chunks = []
    metadata = {"page_count": 0, "total_chars": 0}

    doc = fitz.open(file_path)
    metadata["page_count"] = len(doc)

    for page_num, page in enumerate(doc, 1):
        text = page.get_text()
        text = text.strip()

        if text:  # Skip empty pages
            chunks.append(text)
            metadata["total_chars"] += len(text)

    doc.close()

    return chunks, metadata


def parse_docx(file_path: str) -> Tuple[List[str], Dict[str, Any]]:
    """
    Parse a DOCX file into text chunks (by paragraphs).

    Returns:
        Tuple of (list of text chunks, metadata dict)
    """
    chunks = []
    metadata = {
        "page_count": None,  # DOCX doesn't have clear page breaks
        "total_chars": 0,
        "paragraph_count": 0,
    }

    doc = docx.Document(file_path)
    current_chunk = []
    chunk_size = 1000  # Characters per chunk

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        metadata["paragraph_count"] += 1
        current_chunk.append(text)

        # Join and check size
        chunk_text = "\n".join(current_chunk)
        if len(chunk_text) >= chunk_size:
            chunks.append(chunk_text)
            metadata["total_chars"] += len(chunk_text)
            current_chunk = []

    # Add remaining text
    if current_chunk:
        chunk_text = "\n".join(current_chunk)
        if chunk_text.strip():
            chunks.append(chunk_text)
            metadata["total_chars"] += len(chunk_text)

    return chunks, metadata


def parse_txt(file_path: str) -> Tuple[List[str], Dict[str, Any]]:
    """
    Parse a TXT file into text chunks.

    Returns:
        Tuple of (list of text chunks, metadata dict)
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    metadata = {"page_count": None, "total_chars": len(content)}

    # Split by paragraphs or lines
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

    # If no clear paragraphs, split by lines
    if not paragraphs:
        paragraphs = [line.strip() for line in content.split("\n") if line.strip()]

    # Group into chunks
    chunks = []
    current_chunk = []
    chunk_size = 1000

    for para in paragraphs:
        current_chunk.append(para)
        chunk_text = "\n\n".join(current_chunk)

        if len(chunk_text) >= chunk_size:
            chunks.append(chunk_text)
            current_chunk = []

    if current_chunk:
        chunk_text = "\n\n".join(current_chunk)
        if chunk_text.strip():
            chunks.append(chunk_text)

    return chunks, metadata


def parse_document(file_path: str) -> Tuple[List[str], Dict[str, Any]]:
    """
    Parse any supported document type.

    Args:
        file_path: Path to the document file

    Returns:
        Tuple of (list of text chunks, metadata dict)
    """
    ext = Path(file_path).suffix.lower()

    parsers = {".pdf": parse_pdf, ".docx": parse_docx, ".txt": parse_txt}

    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"Unsupported file type: {ext}")

    return parser(file_path)


def create_llama_documents(
    chunks: List[str], filename: str, doc_id: str, category: Optional[str] = None
) -> List[LlamaDocument]:
    """
    Convert text chunks to LlamaIndex Document objects with metadata.
    """
    documents = []

    for i, chunk in enumerate(chunks, 1):
        doc = LlamaDocument(
            text=chunk,
            metadata={
                "doc_id": doc_id,
                "filename": filename,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "category": category,
            },
        )
        documents.append(doc)

    return documents
